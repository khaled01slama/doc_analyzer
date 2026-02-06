"""Document parsing for PDF and Word files."""

import io
from pathlib import Path
from docx import Document as DocxDocument
import pdfplumber


async def parse_document(content: bytes, filename: str) -> tuple[str, int]:
    """Parse PDF or Word document and extract text."""
    extension = Path(filename).suffix.lower()

    if extension == ".pdf":
        return await _parse_pdf(content)
    elif extension in [".docx", ".doc"]:
        return await _parse_docx(content)
    else:
        raise ValueError(f"Unsupported file type: {extension}")


async def _parse_pdf(content: bytes) -> tuple[str, int]:
    """Extract text from PDF."""
    text_parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(page_text)
    return "\n\n".join(text_parts), page_count


async def _parse_docx(content: bytes) -> tuple[str, int]:
    """Extract text from Word document."""
    doc = DocxDocument(io.BytesIO(content))
    text_parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                text_parts.append(row_text)

    full_text = "\n\n".join(text_parts)
    estimated_pages = max(1, len(full_text) // 3000)
    return full_text, estimated_pages
